import random
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.shared import Pt  # 用于设置间距
import tkinter as tk
from tkinter import messagebox, filedialog  # 新增文件对话框


def generate_question(selected_operations):
    operators = selected_operations
    num1 = random.randint(1, 100)
    num2 = random.randint(1, 100)
    operator = random.choice(operators)
    
    # 确保生成的题目符合条件
    if operator == '+':
        return num1, operator, num2, num1 + num2
    elif operator == '-':
        while num2 > num1:  # 保证减法不为负数
            num1 ,num2 = num2 ,num1
        return num1, operator, num2, num1 - num2
    elif operator == '*':
        return num1, '×', num2, num1 * num2
    elif operator == '/':
        num2 = random.randint(1, 10)
        num1 = num2 * random.randint(1, 20)  # 确保能整除
        return num1, '÷', num2, num1

# 生成题目的函数
def generate_questions(quantity, selected_operations):
    questions = []
    answers = []
    
    for i in range(quantity):
        num1, operator, num2, answer = generate_question(selected_operations)
        question = f"{num1} {operator} {num2} = " 
        questions.append(question)
        answers.append(f"{question} {answer}")
    # print(len(questions),len(answers))
    return questions, answers

def save_to_word(questions, answers):
    filename = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
    )
    if not filename:
        return False  # 用户取消选择
    
    try:
        doc = Document()
        
        # 题目标题
        title = doc.add_heading('数学题目', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 题目分栏设置 插入"连续分节"，设置4栏，题目区域都属于这一节
        section = doc.add_section(start_type=WD_SECTION_START.CONTINUOUS)
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set('num', '4')

        # 添加题目内容
        for q in questions:
            p = doc.add_paragraph(q)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 左对齐
            p.paragraph_format.space_after = Pt(9)

        # 分页
        doc.add_page_break()

        # 答案标题
        answer_title = doc.add_heading('答案', 0)
        answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 答案分栏设置
        section = doc.add_section(start_type=WD_SECTION_START.CONTINUOUS)
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set('num', '4')

        # 添加答案内容
        for a in answers:
            p = doc.add_paragraph(a)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(9)

        doc.save(filename)
        return True  # 保存成功
    except Exception as e:
        messagebox.showerror("保存失败", f"生成文档时出错：{str(e)}")
        return False

# GUI界面
def on_generate():
    try:
        quantity = int(entry_quantity.get())
        selected_operations = []

        if quantity < 1 or quantity > 500:
            messagebox.showerror("输入错误", "题目数量请控制在1-500之间")
            return
        
        if var_add.get():
            selected_operations.append('+')
        if var_sub.get():
            selected_operations.append('-')
        if var_mul.get():
            selected_operations.append('*')
        if var_div.get():
            selected_operations.append('/')
        
        if not selected_operations:
            selected_operations = ['+', '-', '*', '/']
        
        questions, answers = generate_questions(quantity, selected_operations)
        # 保存文档并检查结果
        if save_to_word(questions, answers):
            messagebox.showinfo("生成成功", f"已成功生成 {quantity} 道题目及答案！")
        
    except ValueError:
        messagebox.showerror("输入错误", "请输入有效的题目数量（1-500）")

# 创建主窗口
root = tk.Tk()
root.title("小学生出题工具")
root.geometry("400x350")  
root.config(bg="#f5f5f5")  

label_quantity = tk.Label(root, text="请输入题目数量:", font=("微软雅黑", 12), bg="#f5f5f5")
label_quantity.grid(row=0, column=0, padx=20, pady=10, sticky="w")

entry_quantity = tk.Entry(root, font=("微软雅黑", 12), width=15)
entry_quantity.insert(0, "100")
entry_quantity.grid(row=0, column=1, padx=20, pady=10)

label_operations = tk.Label(root, text="请选择题目类型:", font=("微软雅黑", 12), bg="#f5f5f5")
label_operations.grid(row=1, column=0, padx=20, pady=10, sticky="w")

var_add = tk.BooleanVar()
var_sub = tk.BooleanVar()
var_mul = tk.BooleanVar()
var_div = tk.BooleanVar()

check_add = tk.Checkbutton(root, text="加法", variable=var_add, font=("微软雅黑", 12))
check_add.grid(row=2, column=0, padx=20, sticky="w")

check_sub = tk.Checkbutton(root, text="减法", variable=var_sub, font=("微软雅黑", 12))
check_sub.grid(row=3, column=0, padx=20, sticky="w")

check_mul = tk.Checkbutton(root, text="乘法", variable=var_mul, font=("微软雅黑", 12))
check_mul.grid(row=4, column=0, padx=20, sticky="w")

check_div = tk.Checkbutton(root, text="除法", variable=var_div, font=("微软雅黑", 12))
check_div.grid(row=5, column=0, padx=20, sticky="w")

button_generate = tk.Button(root, text="生成题目", command=on_generate, font=("微软雅黑", 14), bg="#4CAF50", fg="white")
button_generate.grid(row=6, column=0, columnspan=2, pady=40)

root.mainloop()

